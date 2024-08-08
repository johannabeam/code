import pandas as pd
!pip install python-docx
from docx import Document
from docx.shared import Pt


# Function to format and combine column values into a single block of text
def format_row_cells(row):
    return (
        f"{row['Species_code']} ({row['Sample Type']})\n"
        f"{row['Band_no']}\n"
        f"Ext: {row['Date']}\n"
        f"{row['CityTown']}, {row['State']} {row['Country_code']}"
    )

# Load CSV file with different encodings ## THIS IS WHERE YOU MATCH THE INPUT FILE NAMES! ##
csv_file = 'LOSH_Blood_Stickers.csv'
encodings = ['utf-8', 'ISO-8859-1', 'cp1252']  # List of potential encodings

df = None
for encoding in encodings:
    try:
        df = pd.read_csv(csv_file, encoding=encoding)
        print(f"Successfully read CSV with encoding: {encoding}")
        break
    except (FileNotFoundError, pd.errors.EmptyDataError, pd.errors.ParserError) as e:
        print(f"Error reading CSV with encoding {encoding}: {e}")
    except Exception as e:
        print(f"Error: {e}")

if df is None:
    print("Failed to read the CSV file. Exiting.")
    exit()

# Inspect and clean column names
print("Original column names:", df.columns.tolist())
df.columns = df.columns.str.strip()  # Remove leading/trailing spaces
print("Cleaned column names:", df.columns.tolist())

# Load the existing Word document ## THIS IS WHERE YOU MATCH THE INPUT FILE NAMES! ##
doc_path = 'side_label.docx'
doc = None
table = None

try:
    doc = Document(doc_path)
    if doc.tables:
        table = doc.tables[0]
    else:
        raise Exception("No tables found in the document.")
except FileNotFoundError:
    print(f"Error: The file '{doc_path}' does not exist.")
    exit()
except Exception as e:
    print(f"Error loading Word document or table: {e}")
    exit()

# Flatten the table into a list of cells
table_cells = [cell for row in table.rows for cell in row.cells]

# Flatten the DataFrame into a list of formatted text
formatted_texts = [format_row_cells(row) for _, row in df.iterrows()]

# Update each cell in the table
for cell, text in zip(table_cells, formatted_texts):
    # Clear existing text in the cell
    for paragraph in cell.paragraphs:
        cell._element.remove(paragraph._element)

    # Add the formatted text to the cell
    p = cell.add_paragraph()
    run = p.add_run(text)

    # Set font size
    run.font.size = Pt(5.5)  # Set font size to 5.5 points

# Save the modified Word document
try:
    doc.save('LOSH_side_labels.docx')
except Exception as e:
    print(f"Error saving Word document: {e}")


# FOR MAKING THE TOP LABELS, WHICH HAVE A SPECIAL FORMAT OF THE BGP ID #   
import pandas as pd
!pip install python-docx
from docx import Document
from docx.shared import Pt


# Function to format and combine column values into a single block of text
def format_row_cells(row):
    # Split the BGP_ID into parts based on the delimiter 'N'
    bgp_id_parts = row['BGP_ID'].split('N')

    # Handle cases where 'N' is not found or there are more than two parts
    if len(bgp_id_parts) == 2:
        bgp_id_first_part = bgp_id_parts[0] + 'N'  # Add 'N' back to the first part
        bgp_id_second_part = bgp_id_parts[1]
    else:
        # If the delimiter 'N' is not found or there are more than 2 parts, handle accordingly
        bgp_id_first_part = row['BGP_ID']  # Keep it as-is if 'N' is not found
        bgp_id_second_part = ''

    # Return the formatted block of text
    return (
        f"{row['Species_code']}\n"
        f"{bgp_id_first_part}\n"
        f"{bgp_id_second_part}"
    )
# Load CSV file with different encodings
csv_file = 'LOSH_Blood_Stickers.csv'
encodings = ['utf-8', 'ISO-8859-1', 'cp1252']  # List of potential encodings

df = None
for encoding in encodings:
    try:
        df = pd.read_csv(csv_file, encoding=encoding)
        print(f"Successfully read CSV with encoding: {encoding}")
        break
    except (FileNotFoundError, pd.errors.EmptyDataError, pd.errors.ParserError) as e:
        print(f"Error reading CSV with encoding {encoding}: {e}")
    except Exception as e:
        print(f"Error: {e}")

if df is None:
    print("Failed to read the CSV file. Exiting.")
    exit()

# Inspect and clean column names
print("Original column names:", df.columns.tolist())
df.columns = df.columns.str.strip()  # Remove leading/trailing spaces
print("Cleaned column names:", df.columns.tolist())

# Load the existing Word document
doc_path = 'top_label.docx'
doc = None
table = None

try:
    doc = Document(doc_path)
    if doc.tables:
        table = doc.tables[0]
    else:
        raise Exception("No tables found in the document.")
except FileNotFoundError:
    print(f"Error: The file '{doc_path}' does not exist.")
    exit()
except Exception as e:
    print(f"Error loading Word document or table: {e}")
    exit()

# Flatten the table into a list of cells
table_cells = [cell for row in table.rows for cell in row.cells]

# Flatten the DataFrame into a list of formatted text
formatted_texts = [format_row_cells(row) for _, row in df.iterrows()]

# Update each cell in the table
for cell, text in zip(table_cells, formatted_texts):
    # Clear existing text in the cell
    for paragraph in cell.paragraphs:
        cell._element.remove(paragraph._element)

    # Add the formatted text to the cell
    p = cell.add_paragraph()
    run = p.add_run(text)

    # Set font size
    run.font.size = Pt(5.5)  # Set font size to 5.5 points

# Save the modified Word document
try:
    doc.save('finished_top_label.docx')
except Exception as e:
    print(f"Error saving Word document: {e}")




# AND FINALLY, FOR MAKING THE TISSUE/BLOOD TUBE LABELS, WHICH HAVE A FUNKY FORMATTED TEMPLATE. #

import pandas as pd
!pip install python-docx
from docx import Document
from docx.shared import Pt

import pandas as pd
from docx import Document
from docx.shared import Pt

# Function to format and combine column values into a single block of text
def format_row_cells(row):
    # Split the BGP_ID into parts based on the delimiter 'N'
    bgp_id_parts = row['BGP_ID'].split('N')

    # Handle cases where 'N' is not found or there are more than two parts
    if len(bgp_id_parts) == 2:
        bgp_id_first_part = bgp_id_parts[0] + 'N'  # Add 'N' back to the first part
        bgp_id_second_part = bgp_id_parts[1]
    else:
        # If the delimiter 'N' is not found or there are more than 2 parts, handle accordingly
        bgp_id_first_part = row['BGP_ID']  # Keep it as-is if 'N' is not found
        bgp_id_second_part = ''

    # Return the formatted block of text
    return (
        f"\n"
        f"{bgp_id_first_part}\n"
        f"{bgp_id_second_part}"
    )

# Load CSV file with different encodings
csv_file = 'YEWA_blood.csv'
encodings = ['utf-8', 'ISO-8859-1', 'cp1252']  # List of potential encodings

df = None
for encoding in encodings:
    try:
        df = pd.read_csv(csv_file, encoding=encoding)
        print(f"Successfully read CSV with encoding: {encoding}")
        break
    except (FileNotFoundError, pd.errors.EmptyDataError, pd.errors.ParserError) as e:
        print(f"Error reading CSV with encoding {encoding}: {e}")
    except Exception as e:
        print(f"Error: {e}")

if df is None:
    print("Failed to read the CSV file. Exiting.")
    exit()

# Inspect and clean column names
print("Original column names:", df.columns.tolist())
df.columns = df.columns.str.strip()  # Remove leading/trailing spaces
print("Cleaned column names:", df.columns.tolist())

# Load the existing Word document
doc_path = 'tissue_label.docx'
doc = None
table = None

try:
    doc = Document(doc_path)
    if doc.tables:
        table = doc.tables[0]
    else:
        raise Exception("No tables found in the document.")
except FileNotFoundError:
    print(f"Error: The file '{doc_path}' does not exist.")
    exit()
except Exception as e:
    print(f"Error loading Word document or table: {e}")
    exit()

# Flatten the DataFrame into a list of formatted text
formatted_texts = [format_row_cells(row) for _, row in df.iterrows()]

# Get table dimensions
num_rows = len(table.rows)
num_cols = len(table.columns)

# Ensure that we have enough text to fill the table
if len(formatted_texts) > (num_rows // 2 + (num_rows % 2)) * (num_cols // 2 + (num_cols % 2)):
    raise Exception("Not enough space in the table for all formatted texts")

text_index = 0
for row_idx in range(num_rows):
    if row_idx % 2 == 0:  # Skip rows to print only every other row
        for col_idx in range(0, num_cols, 2):  # Print into every other cell
            cell_idx = row_idx * num_cols + col_idx
            if text_index < len(formatted_texts):
                cell = table.cell(row_idx, col_idx)
                text = formatted_texts[text_index]
                text_index += 1

                # Clear existing text in the cell
                for paragraph in cell.paragraphs:
                    cell._element.remove(paragraph._element)

                # Add the formatted text to the cell
                p = cell.add_paragraph()
                run = p.add_run(text)

                # Set font size
                run.font.size = Pt(5.5)  # Set font size to 5.5 points

# Save the modified Word document
try:
    doc.save('finished_YEWA_label.docx')
except Exception as e:
    print(f"Error saving Word document: {e}")
